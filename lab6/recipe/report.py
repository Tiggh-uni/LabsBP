# recipe_package/report.py
import openpyxl
from docx import Document

def save_to_xls(results, filepath):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Результаты"
    ws.append(["Рецепт", "Ккал", "Белки", "Жиры", "Углеводы", "Стоимость (руб)"])
    ws.append([
        results["recipe"],
        results["kcal"],
        results["protein"],
        results["fat"],
        results["carbs"],
        results["price"]
    ])
    # Детализация по ингредиентам
    ws2 = wb.create_sheet("Ингредиенты")
    ws2.append(["Ингредиент", "Грамм", "Ккал", "Цена (руб)"])
    for ing, (grams, kcal, price) in results["breakdown"].items():
        ws2.append([ing, grams, kcal, price])
    wb.save(filepath)

def save_to_doc(results, filepath):
    doc = Document()
    doc.add_heading('Отчёт по рецепту', level=1)
    doc.add_paragraph(f'Блюдо: {results["recipe"]}')
    doc.add_paragraph(f'Энергетическая ценность: {results["kcal"]} ккал')
    p = doc.add_paragraph()
    p.add_run('Белки: ').bold = True
    p.add_run(f'{results["protein"]} г, ')
    p.add_run('Жиры: ').bold = True
    p.add_run(f'{results["fat"]} г, ')
    p.add_run('Углеводы: ').bold = True
    p.add_run(f'{results["carbs"]} г')
    doc.add_paragraph(f'Стоимость: {results["price"]} руб.')
    doc.add_heading('Состав:', level=2)
    for ing, (grams, kcal, price) in results["breakdown"].items():
        doc.add_paragraph(f'{ing}: {grams} г, {kcal:.1f} ккал, {price:.2f} руб.')
    doc.save(filepath)